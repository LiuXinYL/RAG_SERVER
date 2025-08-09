import os
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def convert_md_to_word(md_path, output_file):
    """
    将 Markdown 内容转换为 Word 文档

    参数:
    md_content (str): Markdown 文本内容
    output_file (str): 输出 Word 文件路径
    """

    # 读取Markdown文件内容
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 将 Markdown 转换为 HTML
    html = markdown(md_text, extensions=['markdown.extensions.fenced_code',
                                            'markdown.extensions.tables',
                                            'markdown.extensions.nl2br'])
    # 使用 BeautifulSoup 解析 HTML
    soup = BeautifulSoup(html, 'html.parser')

    # 保存解析后的 HTML 到文件
    html_output_path = os.path.splitext(output_file)[0] + '_parsed.html'
    with open(html_output_path, 'w', encoding='utf-8') as html_file:
        html_file.write(soup.prettify())


    # 创建 Word 文档
    doc = Document()

    # 设置中文字体支持
    doc.styles['Normal'].font.name = 'SimHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    # 处理 HTML 元素并添加到 Word 文档
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                                  'ul', 'ol', 'li', 'pre', 'code', 'table', 'tr', 'td', 'th']):
        if element.name.startswith('h'):
            # 处理标题
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)

        elif element.name == 'p':
            # 处理段落
            p = doc.add_paragraph()
            p.add_run(element.get_text())

        elif element.name == 'ul':
            # 处理无序列表
            pass  # 列表项会在 li 标签中处理

        elif element.name == 'ol':
            # 处理有序列表
            pass  # 列表项会在 li 标签中处理

        elif element.name == 'li':
            # 处理列表项
            if element.parent.name == 'ul':
                doc.add_paragraph(element.get_text(), style='List Bullet')
            elif element.parent.name == 'ol':
                doc.add_paragraph(element.get_text(), style='List Number')

        elif element.name == 'pre':
            # 处理代码块
            code_block = element.get_text()
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(code_block)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)

        elif element.name == 'table':
            # 处理表格
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text()

    # 保存 Word 文档
    doc.save(output_file)
    print(f"已成功将 Markdown 转换为 Word 文档: {output_file}")



if __name__ == "__main__":

    md_path = 'C:/Users/16000/Desktop/ww.md'
    word_path = 'C:/Users/16000/Desktop/wwq1.docx'

    # 标题一 宋体 三号
    # 标题二 宋体 四号
    # 标题三 宋体 小四
    # 标题四 宋体 五号
    # 正文中文 宋体 五号
    # 正文英文 Calibri 五号

    # 执行转换
    convert_md_to_word(md_path, word_path)

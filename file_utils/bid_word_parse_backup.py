import pandas as pd
import numpy as np
from docx import Document
from docx.shared import RGBColor
from docx.table import Table
from find_headline import identify_and_check_captions, format_result

pd.set_option('display.max_rows', 200)  # 最大显示行数
pd.set_option('display.max_columns', 50)  # 最大显示列数
pd.set_option('display.width', 1000)  # 显示宽度
pd.set_option('display.max_colwidth', 200)  # 每列最大宽度


def _is_heading(paragraph):
    """判断段落是否为标题及标题级别"""
    if paragraph.style.name.startswith('Heading'):
        try:
            level = int(paragraph.style.name.split(' ')[-1])
            return True, level
        except:
            return True, "Unable to parse"  # 无法解析级别时默认为1
    return False, "Not Heading"


def _find_table_after_heading(doc, start_para_idx):
    """在指定段落索引之后查找第一个表格"""
    # 搜索范围限制在标题后的N个元素内，避免遍历整个文档
    SEARCH_LIMIT = 20  # 可根据实际文档结构调整

    # 从标题段落之后的元素开始查找
    for i in range(start_para_idx + 1, min(start_para_idx + SEARCH_LIMIT + 1, len(doc.element.body))):
        element = doc.element.body[i]

        # 直接检查元素是否为表格
        if element.tag.endswith('tbl'):
            return Table(element, doc)

        # 处理可能包含表格的其他元素类型（如单元格）
        # 这里我们检查元素的所有子元素，看是否包含表格
        for child in element.iterchildren():
            if child.tag.endswith('tbl'):
                return Table(child, doc)

    # 未找到表格
    return None


def _get_line_spacing(paragraph):
    """获取段落的行间距信息"""
    if paragraph.paragraph_format.line_spacing_rule is None:
        return "未设置"
    elif paragraph.paragraph_format.line_spacing_rule == 1:  # 单倍行距
        return "单倍行距"
    elif paragraph.paragraph_format.line_spacing_rule == 2:  # 1.5 倍行距
        return "1.5 倍行距"
    elif paragraph.paragraph_format.line_spacing_rule == 3:  # 双倍行距
        return "双倍行距"
    elif paragraph.paragraph_format.line_spacing_rule == 4:  # 固定值
        return f"固定值: {paragraph.paragraph_format.line_spacing.pt} 磅"
    elif paragraph.paragraph_format.line_spacing_rule == 5:  # 多倍行距
        return f"多倍行距: {paragraph.paragraph_format.line_spacing} 倍"
    else:
        return "未知格式"


def save_file(font_df, table_dict, note_df, save_path):
    # 创建一个ExcelWriter对象
    with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
        # 将字体信息写入第一个工作表
        font_df.to_excel(writer, sheet_name='word_parse', index=False)

        # 将“注”信息写入第二个工作表
        note_df.to_excel(writer, sheet_name='note_info', index=False)

        # 将每个表格的数据写入单独的工作表
        for table_name, table_df in table_dict.items():
            # 使用表格标题作为工作表名称，确保不超过31个字符

            table_df.reset_index(drop=True, inplace=True)
            table_df.to_excel(writer, sheet_name=table_name, index=False)

    print(f"已将所有结果写入到 {save_path}")


def _process_paragraph(paragraph, context, is_heading, heading_level, line_spacing, is_table_content=False):
    """处理单个段落，提取所有 Run 的格式信息"""
    runs_info = []
    for run in paragraph.runs:
        text = run.text.strip()
        if not text:
            continue

        font_name = run.font.name if run.font.name else "默认字体"
        font_size = run.font.size.pt if run.font.size else "未找到"

        font_color = "未设置"
        if run.font.color and isinstance(run.font.color.rgb, RGBColor):
            rgb = run.font.color.rgb
            font_color = f"RGB({rgb[0]}, {rgb[1]}, {rgb[2]})"

        runs_info.append({
            "text": text,
            "font_name": font_name,
            "font_size": font_size,
            "font_color": font_color,
            "context": context,
            "is_heading": is_heading,
            "heading_level": heading_level,
            "line_spacing": line_spacing,
            "is_table_content": is_table_content  # 添加表格内容标记
        })
    return runs_info


def get_font_info(doc):
    font_info_list = []
    special_tables = {"开标一览表": None, "投标分项报价表": None}  # 存储特殊表格信息
    missing_tables = {"开标一览表": False, "投标分项报价表": False}  # 标记未找到的表格
    note_info_list = []  # 存储“注”信息

    # 1. 遍历所有段落，定位特殊标题并查找其后的表格
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        if para_text:
            context = f"正文 {para_idx + 1}"
            is_heading, heading_level = _is_heading(paragraph)
            line_spacing = _get_line_spacing(paragraph)

            # 处理段落并获取字体信息
            runs_info = _process_paragraph(
                paragraph, context, is_heading, heading_level, line_spacing, False
            )

            # 检查是否为“注”行
            if para_text[0] == '注':
                note_info_list.extend(runs_info)
            else:
                font_info_list.extend(runs_info)

            # 检查是否为特殊表格标题
            table_title = paragraph.text.strip()
            if is_heading and (heading_level == 1 or heading_level == 2 or heading_level == 3):
                for keyword in special_tables:
                    if keyword in table_title:
                        # 找到匹配的关键词，开始搜索其后的表格
                        table = _find_table_after_heading(doc, para_idx)
                        if table:
                            # 处理表格内容
                            for row_idx, row in enumerate(table.rows):
                                for cell_idx, cell in enumerate(row.cells):
                                    context = f"{table_title}_行{row_idx + 1}_列{cell_idx + 1}"
                                    for para_in_cell in cell.paragraphs:
                                        if para_in_cell.text.strip():
                                            is_h, h_level = _is_heading(para_in_cell)
                                            ls = _get_line_spacing(para_in_cell)
                                            # 处理表格内的段落，标记为表格内容
                                            font_info_list.extend(
                                                _process_paragraph(
                                                    para_in_cell, context, is_h, h_level, ls, True
                                                )
                                            )

                            # 保存表格数据
                            table_df = pd.DataFrame([
                                [cell.text for cell in row.cells] for row in table.rows
                            ])
                            table_df.attrs['table_title'] = table_title
                            special_tables[keyword] = table_df
                        else:
                            missing_tables[keyword] = True
                        break

    return pd.DataFrame(font_info_list), special_tables, missing_tables, pd.DataFrame(note_info_list)


def check_format_consistency(df, is_heading=True, group_column='heading_level'):
    """
    检查格式的一致性，支持标题和正文检测

    参数:
    df -- 包含文本信息的DataFrame
    is_heading -- 是否为标题数据
    group_column -- 分组依据的列名（仅对标题有效）

    返回:
    inconsistency_report -- 不一致情况报告
    """
    # 确保数据中包含必要的列
    required_columns = ['text', 'font_name', 'font_size', 'line_spacing', 'context']

    if not all(col in df.columns for col in required_columns):
        raise ValueError(f"DataFrame缺少必要的列，请检查数据结构: {required_columns}")

    # 用于存储不一致报告
    inconsistency_report = []

    if is_heading:
        # 标题按级别分组检查
        groups = df.groupby(group_column)

        for level, group in groups:
            # 检查字体一致性
            font_names = group['font_name'].unique()
            if len(font_names) > 1 or np.nan in font_names:
                inconsistency_report.append({
                    'type': 'heading',
                    'level': level,
                    'property': 'font_name',
                    'expected': font_names[0] if not pd.isna(font_names[0]) else None,
                    'actual': font_names,
                    'inconsistent_items': group[~group['font_name'].eq(font_names[0])][
                        ['text', 'font_name', 'context']].to_dict('records')
                })

            # 检查字号一致性
            font_sizes = group['font_size'].unique()
            if len(font_sizes) > 1 or np.nan in font_sizes:
                inconsistency_report.append({
                    'type': 'heading',
                    'level': level,
                    'property': 'font_size',
                    'expected': font_sizes[0] if not pd.isna(font_sizes[0]) else None,
                    'actual': font_sizes,
                    'inconsistent_items': group[~group['font_size'].eq(font_sizes[0])][
                        ['text', 'font_size', 'context']].to_dict('records')
                })

            # 检查行间距一致性
            line_spacings = group['line_spacing'].unique()
            if len(line_spacings) > 1 or np.nan in line_spacings:
                inconsistency_report.append({
                    'type': 'heading',
                    'level': level,
                    'property': 'line_spacing',
                    'expected': line_spacings[0] if not pd.isna(line_spacings[0]) else None,
                    'actual': line_spacings,
                    'inconsistent_items': group[~group['line_spacing'].eq(line_spacings[0])][
                        ['text', 'line_spacing', 'context']].to_dict('records')
                })

        # inconsistency_report.append({"checkName": "标题检测", "result": True, "message": "提示: 未发现标题问题"})

    else:
        # 正文整体检查
        # 检查字体一致性
        font_names = df['font_name'].unique()
        if len(font_names) > 1 or np.nan in font_names:
            inconsistency_report.append({
                'type': 'body',
                'property': 'font_name',
                'expected': font_names[0] if not pd.isna(font_names[0]) else None,
                'actual': font_names,
                'inconsistent_items': df[~df['font_name'].eq(font_names[0])][['text', 'font_name', 'context']].to_dict(
                    'records')
            })

        # 检查字号一致性
        font_sizes = df['font_size'].unique()
        if len(font_sizes) > 1 or np.nan in font_sizes:
            inconsistency_report.append({
                'type': 'body',
                'property': 'font_size',
                'expected': font_sizes[0] if not pd.isna(font_sizes[0]) else None,
                'actual': font_sizes,
                'inconsistent_items': df[~df['font_size'].eq(font_sizes[0])][['text', 'font_size', 'context']].to_dict(
                    'records')
            })

        # 检查行间距一致性
        line_spacings = df['line_spacing'].unique()
        if len(line_spacings) > 1 or np.nan in line_spacings:
            inconsistency_report.append({
                'type': 'body',
                'property': 'line_spacing',
                'expected': line_spacings[0] if not pd.isna(line_spacings[0]) else None,
                'actual': line_spacings,
                'inconsistent_items': df[~df['line_spacing'].eq(line_spacings[0])][
                    ['text', 'line_spacing', 'context']].to_dict('records')
            })



    return inconsistency_report


def report_info(report):
    # 定义英文属性名到中文的映射
    property_mapping = {
        'line_spacing': '行间距',
        'font_size': '字号',
        'font_name': '字体'
    }

    if not report:
        return {"checkName": "未发现问题", "result": True, "message": f"所有格式检查通过，未发现不一致情况。"}

    heading_issues = [issue for issue in report if issue['type'] == 'heading']
    body_issues = [issue for issue in report if issue['type'] == 'body']

    check_result = []
    # 打印标题问题
    if heading_issues:
        for issue in heading_issues:
            level = issue['level']
            prop = issue['property']
            prop_cn = property_mapping.get(prop, prop)  # 获取中文属性名
            expected = issue['expected']
            actual = issue['actual']

            result_heading = {
                "checkName": f"标题级别 {level} 存在 {prop_cn} 不一致问题",
                "result": False
            }

            inconsistent_texts = []
            for item in issue['inconsistent_items']:
                value = item[prop]
                text = item['text']
                context = item['context']
                inconsistent_texts.append(f"位置：{context}, 内容为：{text}")

            result_heading["message"] = "\n".join(inconsistent_texts)
            check_result.append(result_heading)

    # 打印正文问题
    if body_issues:
        for issue in body_issues:
            prop = issue['property']
            prop_cn = property_mapping.get(prop, prop)  # 获取中文属性名
            expected = issue['expected']
            actual = issue['actual']

            inconsistent_texts = []
            note_flag = False
            for item in issue['inconsistent_items']:
                value = item[prop]
                text = item['text']
                if "注" in text:
                    note_flag = True
                context = item['context']
                inconsistent_texts.append(f"所在位置：{context}, 内容为：{text}")

            if note_flag:
                title = "注"
            else:
                title = "正文"

            result_body = {
                "checkName": f"{title} 存在 {prop_cn} 不一致问题",
                "result": False,
                "message": "\n".join(inconsistent_texts)
            }

            check_result.append(result_body)

    return check_result

def work_flow(doc_, save_path=None):

    font_df, special_tables, missing_tables, note_df = get_font_info(doc_)

    if save_path is not None:
        save_file(font_df, special_tables, note_df, save_path)  # 保存文件到excel, 附加表格信息

    heading_report = []
    body_report = []
    table_report = []
    note_report = []
    # 检测图号、表号是否连续
    result = identify_and_check_captions(doc_)
    captions_report = format_result(result)

    # 检查是否有未找到的表格
    for table_title, missing in missing_tables.items():
        if missing or special_tables[table_title] is None:
            table_report.append(
                {"checkName": f"{table_title}检测", "result": False, "message": f"提示: 文档中未找到 '{table_title}'"}
            )
        else:
            table_report.append({"checkName": f"{table_title}检测", "result": True, "message": f"提示: 文档中未找到 '{table_title}'"})

    # 检查是否有未找到的注信息
    if note_df.empty:
        note_report.append({"checkName": "注信息检测", "result": False, "message": f"提示: 文档中未找到 '注'"})
    else:
        note_info = check_format_consistency(note_df, is_heading=False)
        note_report = report_info(note_info)

    # 检查标题和正文的格式一致性
    heading_df = font_df[font_df['is_heading']]
    heading_info = check_format_consistency(heading_df, is_heading=True)
    if len(heading_info) > 0:
        heading_report = (report_info(heading_info))
    else:
        heading_report.append({"checkName": "标题检测", "result": True, "message": "未发现标题异常"})

    body_df = font_df[(~font_df['is_heading']) & (~font_df['is_table_content'])]
    body_info = check_format_consistency(body_df, is_heading=False)
    if len(body_info) > 0:
        body_report = (report_info(body_info))
    else:
        body_report.append({"checkName": "标题检测", "result": True, "message": "未发现标题异常"})

    # 打印报告
    check_report_list = captions_report + table_report + note_report + heading_report + body_report

    return check_report_list


if __name__ == "__main__":

    # docx_path = 'C:/Users/16000/Desktop/小红小明测试.docx'
    docx_path = 'C:/Users/16000/Desktop/投标表格测试.docx'
    excel_save_path = 'C:/Users/16000/Desktop/小红小明测试_综合结果.xlsx'

    doc_ = Document(docx_path)

    check_list = work_flow(doc_, excel_save_path)
    print('check_list', check_list)
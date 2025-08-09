import os
import sys
import xml.etree.ElementTree as ET
from glob import glob
from pathlib import Path
from typing import List

import docx
import numpy as np
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from sklearn.metrics.pairwise import cosine_similarity
from tqdm import tqdm

from llm_init import LLM_INIT

# 获取当前项目目录的绝对路径
current_dir = os.path.abspath(os.path.dirname(__file__))
# 添加到 sys.path
sys.path.append(current_dir)


def check_and_create_dir(dir_path):
    """面向对象的目录检查与创建"""
    path = Path(dir_path)
    try:
        if not path.is_dir():
            path.mkdir(parents=True, exist_ok=True)
            print(f"目录 '{dir_path}' 不存在，已自动创建")
        else:
            print(f"目录 '{dir_path}' 已存在")
    except PermissionError:
        print(f"权限不足，无法创建目录 '{dir_path}'")
    except Exception as e:
        print(f"未知错误: {str(e)}")


def check_xml_validity(xml_path):
    try:
        ET.parse(xml_path)  # 解析XML文件
        with open(xml_path, encoding='utf-8') as f:
            data = f.read()
        return data

    except ET.ParseError as e:
        print(f"XML格式错误：{e}")
        return False


def read_word_text(file_path):
    # 读取文档内容
    doc = docx.Document(file_path)
    # 提取文本并过滤空段落
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    # 拼接文本并添加段落分隔符
    return '\n\n'.join(paragraphs)


def read_md_text(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def split_document_func(text, max_length, overlap, show_progress=False):
    """
    根据字符长度切分文本，并对切分点有覆盖，支持显示进度条
    """
    if not text:
        return []

    text_length = len(text)
    if max_length >= text_length:
        return [text]

    # 计算分段数量
    step = max_length - overlap
    num_segments = max(1, (text_length - overlap + step - 1) // step)

    # 生成所有分段的起始和结束索引
    starts = np.arange(num_segments) * step
    ends = starts + max_length
    # 确保最后一个分段不超出文本长度
    ends = np.minimum(ends, text_length)
    starts = np.maximum(starts, 0)

    # 使用tqdm包装迭代过程以显示进度条
    segments = []
    indices_iter = zip(starts, ends)

    for start, end in indices_iter:
        segments.append(text[start:end])

    return segments


def _load_and_process_file(file_path, read_text_func, max_length=500, overlap=100):
    try:
        # 读取文本内容
        text = read_text_func(file_path)

        # 文本为空时直接返回
        if not text:
            print(f"文件 {file_path} 不包含文本内容")
            return []

        # 切分文本
        segments = split_document_func(text, max_length, overlap)

        # 直接在 for 循环行上添加 tqdm
        documents = [
            Document(
                page_content=segment,
                metadata={"segment_index": i, "file_path": file_path}
            )
            for i, segment in enumerate(segments)
        ]

        return documents

    except Exception as e:
        print(f"解析文件 {file_path} 时出错: {str(e)}")
        return []


def load_xml_file(input_dir: str) -> List[Document]:
    """
    按子文件夹分类处理XML
    Args:
        input_dir (str): xml文件路径
    Returns:
        List[Document]: 包含文档内容的列表
    """

    try:

        documents = []
        # 遍历子文件夹内的XML文件
        for xml_file in glob(os.path.join(input_dir, '*.xml')):

            sub_name = os.path.basename(os.path.normpath(xml_file)).split('.')[0]
            data = check_xml_validity(xml_file)

            if "嵌套" in sub_name:
                parts = sub_name.split("-")
                node1 = parts[0].split("_")[-1]  # 提取第一个“组织”
                node2 = parts[2].split("_")[-1]  # 提取第三个“组织”

                content = " ".join(f"{'嵌套'}: {str(data)}")
                # 创建metadata字典
                metadata = {"node1": node1, "node2": node2, "relation": '嵌套', "file_path": xml_file}

            elif "完整" in sub_name:

                node = sub_name.split("_")[1]
                content = " ".join(f"{'完整'}: {str(data)}")
                # 创建metadata字典
                metadata = {"node": node, "file_path": xml_file}

            else:

                node = sub_name.split("_")[1]
                content = " ".join(f"{'单节点'}: {str(data)}")
                # 创建metadata字典
                metadata = {"node": node, "file_path": xml_file}

            # 创建Document对象
            doc = Document(page_content=content, metadata=metadata)
            documents.append(doc)

        return documents

    except Exception as e:
        print(f"解析Excel文件时出错: {str(e)}")
        return []


def load_excel_file(file_path: str) -> List[Document]:
    """
    使用pandas加载并解析Excel文件

    Args:
        file_path (str): Excel文件路径
    Returns:
        List[Document]: 包含文档内容的列表
    """
    try:
        # 读取Excel文件
        df = pd.read_excel(file_path)

        documents = []
        # 遍历DataFrame的每一行
        for index, row in df.iterrows():
            # 将所有列的内容组合成文本
            data = row[[col for col in ['问题', '答案'] if col in df.columns]]
            content = " ".join(f"{col}: {str(val)}" for col, val in data.items())

            # 创建metadata字典
            describe = row[[col for col in ['序号', '备注'] if col in df.columns]]
            metadata = {col: val for col, val in describe.items()}
            # 添加文件路径到metadata
            metadata["file_path"] = file_path
            # 创建Document对象
            doc = Document(page_content=content, metadata=metadata)
            documents.append(doc)

        return documents

    except Exception as e:
        print(f"解析Excel文件时出错: {str(e)}")
        return []


def load_word_file(file_path: str, max_length=500, overlap=100) -> List[Document]:
    """
    加载并解析word文件，带进度显示
    """

    return _load_and_process_file(file_path, read_word_text, max_length, overlap)


def load_markdown_file(file_path: str, max_length=500, overlap=100) -> List[Document]:
    """
    加载并解析 Markdown 文件，带进度显示
    """

    return _load_and_process_file(file_path, read_md_text, max_length, overlap)


def load_document(file_path, file_type):
    # 加载文档
    # print("正在加载文档...")

    if file_type == ".xml":
        documents = load_xml_file(file_path)
    elif file_type in [".xlsx", ".xls", ".csv"]:
        documents = load_excel_file(file_path)

    elif file_type == ".md":
        documents = load_markdown_file(file_path)
    elif file_type in [".docx", ".doc"]:
        documents = load_word_file(file_path)

    elif file_type == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        documents = loader.load()
    elif file_type == ".pdf":
        loader = PyPDFLoader(file_path)
        documents = loader.load()
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")
    print(f"已加载文档，共{len(documents)}页")

    return documents


def split_perception_and_recursion_func(file_path, file_type, chunk_size=300, chunk_overlap=100):
    """
    使用递归字符分割器进行文本分割
    markdown 文档使用内容感知分块
    """

    documents = load_document(file_path, file_type)
    # 提取 Document 对象列表中的文本内容并拼接成一个字符串
    all_text = "\n".join([doc.page_content for doc in documents])

    print("正在分割文本...")
    if file_type == ".md":
        # 针对 Markdown 文件使用 MarkdownHeaderTextSplitter 进行内容感知分块
        headers_to_split_on = [
            ("#", "Header 1"),
            ("##", "Header 2"),
            ("###", "Header 3"),
            ("####", "Header 4"),
            ("#####", "Header 5"),
            ("######", "Header 6"),
        ]
        md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)

        documents = md_splitter.split_text(all_text)
        # 添加元数据
        custom_metadata = {"file_path": file_path}
        for doc in documents:
            doc.metadata.update(custom_metadata)

        # 使用 split_text 方法拆分文本
        print(f"已分割文本，共{len(documents)}个块")
        return documents

    else:
        # 其他文件类型使用默认分块
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )

        text_chunk = text_splitter.split_text(all_text)

        custom_metadata = {"file_path": file_path}
        documents_list = [Document(page_content=chunk, metadata=custom_metadata) for chunk in text_chunk]

        print(f"已分割文本，共{len(documents_list)}个块")

        return documents_list


def semantic_split(text_chunks, embeddings_model, threshold=0.7):
    """
    基于语义相似度进行文档切分
    threshold: 相似度阈值，低于该值的将被切分为新段落
    """

    # 基于语义相似度进行切分
    print("正在进行建向量和语义切分...")

    # 使用FAISS构建向量存储
    vector_store = FAISS.from_documents(text_chunks, embeddings_model)

    # 获取所有文档的文本内容
    doc_ids = list(vector_store.index_to_docstore_id.values())
    docs = [vector_store.docstore.search(doc_id) for doc_id in doc_ids if vector_store.docstore.search(doc_id)]
    texts = [doc.page_content for doc in docs]

    # 使用 embed_documents 方法获取嵌入向量
    all_embeddings = vector_store.embedding_function.embed_documents(texts)

    similarity_matrix = cosine_similarity(all_embeddings)

    # 初始化切分点
    split_points = [0]
    current_chunk = 0

    # 遍历相似度矩阵
    for i in range(1, len(similarity_matrix)):
        if similarity_matrix[i, current_chunk] < threshold:
            split_points.append(i)
            current_chunk = i

    split_points.append(len(similarity_matrix))

    # 根据切分点分组文档块
    documents = []
    for i in range(len(split_points) - 1):
        start = split_points[i]
        end = split_points[i + 1]
        chunk_doc_ids = doc_ids[start:end]
        chunk_docs = [vector_store.docstore.search(doc_id) for doc_id in chunk_doc_ids if
                      vector_store.docstore.search(doc_id)]

        content = "\n\n".join([doc.page_content for doc in chunk_docs])
        metadata = {}

        # 创建Document对象
        doc = Document(page_content=content, metadata=metadata)

        documents.append(doc)

    print(f"已完成语义切分，共切分为{len(documents)}个语义段落")
    return documents


def parent_child_split(file_path, file_type, embedding_model, smaller_chunk_size=200, smaller_overlap=40,
                       larger_chunk_size=1000, larger_chunk_overlap=200):
    """
    父子切分逻辑，将文本块切分为大块和小块，并建立映射关系
    """

    documents = load_document(file_path, file_type)

    # 父子切分
    # print("正在进行父子切分...")
    # 初始化大尺寸分割器
    large_text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=larger_chunk_size,
        chunk_overlap=larger_chunk_overlap,
        separators=["\n\n", "\n", " ", ""]
    )

    # 初始化小尺寸分割器
    small_text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=smaller_chunk_size,
        chunk_overlap=smaller_overlap,
        separators=["\n\n", "\n", " ", ""]
    )

    larger_documents_list = []
    smaller_documents_list = []
    small_to_large_mapping = {}

    for doc in documents:
        # 进行大块切分
        large_text_chunks = large_text_splitter.split_text(doc.page_content)
        large_doc_chunk = [Document(page_content=chunk, metadata=doc.metadata) for chunk in large_text_chunks]
        larger_documents_list.extend(large_doc_chunk)

    # 为父分割构建 FAISS 向量存储
    parent_vector_store = FAISS.from_documents(larger_documents_list, embedding_model)

    parent_docstore_ids = list(parent_vector_store.index_to_docstore_id.values())

    for i, doc_id in enumerate(parent_docstore_ids):

        large_doc = parent_vector_store.docstore.search(doc_id)
        if large_doc is None:
            print(f"未找到 ID 为 {doc_id} 的文档")
            continue

        # 对每个大块进行小块切分
        small_text_chunks = small_text_splitter.split_text(large_doc.page_content)
        small_doc_chunk = [Document(page_content=chunk, metadata=large_doc.metadata) for chunk in small_text_chunks]
        smaller_documents_list.extend(small_doc_chunk)

        # 为子分割构建 FAISS 向量存储
        child_vector_store = FAISS.from_documents(small_doc_chunk, embedding_model)
        # 将子向量存储合并到父向量存储
        parent_vector_store.merge_from(child_vector_store)

        child_docstore_ids = list(child_vector_store.index_to_docstore_id.values())

        # 建立小块到大块的映射关系
        for j, child_id in enumerate(child_docstore_ids):
            small_to_large_mapping[child_id] = doc_id

    print(f"已完成父子切分，共生成 {len(smaller_documents_list)} 个小块和 {len(larger_documents_list)} 个大块")
    return smaller_documents_list, larger_documents_list, parent_vector_store, small_to_large_mapping


def process_document(file_path, embedding_model, split_type='simple', max_length=500, overlap=100,
                     large_chunk_size=1000, large_chunk_overlap=200, chunk_threshold=0.8):
    """
    流程控制函数，根据 split_type 选择对应的切分函数

    Args:
        embedding_model: 嵌入模型
        file_path: 文件路径
        split_type: 切分类型，可选值为 'parent_child', 'semantic'
        max_length: 文本切分的最大长度
        overlap: 文本切分的重叠长度
        large_chunk_size: 父子切分中大块的大小
        large_chunk_overlap: 父子切分中大块的重叠长度
        chunk_threshold: 语义切分的相似度阈值
    Returns:
        根据切分类型返回不同的结果
    """
    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        if split_type == 'parent_child':
            return [], [], {}
        return []

    try:
        file_type = os.path.splitext(file_path)[1].lower()

        if split_type == 'simple':
            documents = load_document(file_path, file_type)

        elif split_type == 'perception' and file_type == ".md":  # 针对markdown文件的内容感知
            documents = split_perception_and_recursion_func(
                file_path, file_type, chunk_size=max_length, chunk_overlap=overlap
            )

        elif split_type == 'recursion':
            documents = split_perception_and_recursion_func(
                file_path, file_type, chunk_size=max_length, chunk_overlap=overlap
            )

        elif split_type == 'semantic':
            documents = semantic_split(file_path, embedding_model, threshold=chunk_threshold)

        elif split_type == 'parent_child':
            smaller_documents, larger_documents, parent_vector_store, small_to_large_mapping = parent_child_split(
                file_path, file_type, embedding_model
            )

            return smaller_documents, larger_documents, parent_vector_store, small_to_large_mapping

        else:
            raise ValueError(f"不支持的切分类型: {split_type}, 请检查文档类型: {file_type}")

        return documents

    except Exception as e:
        print(f"处理文档时出错: {str(e)}")
        if split_type == 'parent_child':
            return [], [], {}
        return []


if __name__ == '__main__':
    # data_path = "C:/Users/16000/Desktop/ww.md"
    data_path = "C:/Users/16000/Desktop/ww.docx"

    # data_path = "C:/Users/16000/Desktop/ww.txt"
    # data_path = "C:/Users/16000/Desktop/ww.pdf"
    # data_path = "C:/Users/16000/Desktop/减脂饮食菜谱.xlsx"

    model = LLM_INIT().create_embeddings_client()
    # split_documents = process_document(data_path, model, )
    # split_documents = process_document(data_path, model, split_type='perception')
    # split_documents = process_document(data_path, model, split_type='recursion')
    # print('split_documents', split_documents)

    # 使用父子切分
    smaller_documents_list, larger_documents_list, parent_docstore_ids, small_to_large_mapping= process_document(
        data_path, model, split_type='parent_child'
    )
    print(smaller_documents_list, larger_documents_list, parent_docstore_ids, small_to_large_mapping)

    # # 使用语义切分
    # split_documents = process_document(data_path, model, split_type='semantic')
    # print(split_documents)

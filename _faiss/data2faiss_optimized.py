"""
优化的文档处理和向量化模块
支持多种文档格式的加载、切分和向量化处理
"""
import os
import sys
import xml.etree.ElementTree as ET
from glob import glob
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union, Any
from dataclasses import dataclass
from enum import Enum
import logging

import numpy as np
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from sklearn.metrics.pairwise import cosine_similarity
from tqdm import tqdm
import docx

# 添加当前目录到系统路径
current_dir = os.path.abspath(os.path.dirname(__file__))
sys.path.append(current_dir)

from llm_init import LLM_INIT

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class SplitType(Enum):
    """文档切分类型枚举"""
    SIMPLE = "simple"
    PERCEPTION = "perception"
    RECURSION = "recursion"
    SEMANTIC = "semantic"
    PARENT_CHILD = "parent_child"


class FileType(Enum):
    """支持的文件类型枚举"""
    XML = ".xml"
    EXCEL = (".xlsx", ".xls", ".csv")
    MARKDOWN = ".md"
    WORD = (".docx", ".doc")
    TEXT = ".txt"
    PDF = ".pdf"


@dataclass
class SplitConfig:
    """文档切分配置"""
    chunk_size: int = 500
    chunk_overlap: int = 100
    large_chunk_size: int = 1000
    large_chunk_overlap: int = 200
    semantic_threshold: float = 0.7


class DocumentProcessor:
    """文档处理器基类"""
    
    def __init__(self, config: SplitConfig = None):
        self.config = config or SplitConfig()
        self.supported_extensions = {
            '.xml', '.xlsx', '.xls', '.csv', '.md', 
            '.docx', '.doc', '.txt', '.pdf'
        }
    
    def check_and_create_dir(self, dir_path: Union[str, Path]) -> bool:
        """检查并创建目录"""
        path = Path(dir_path)
        try:
            if not path.is_dir():
                path.mkdir(parents=True, exist_ok=True)
                logger.info(f"目录 '{dir_path}' 已创建")
            return True
        except PermissionError:
            logger.error(f"权限不足，无法创建目录 '{dir_path}'")
            return False
        except Exception as e:
            logger.error(f"创建目录时出错: {str(e)}")
            return False
    
    def validate_file(self, file_path: Union[str, Path]) -> Tuple[bool, Optional[str]]:
        """验证文件是否存在且支持"""
        path = Path(file_path)
        if not path.exists():
            return False, "文件不存在"
        
        file_type = path.suffix.lower()
        if file_type not in self.supported_extensions:
            return False, f"不支持的文件类型: {file_type}"
        
        return True, file_type


class XMLProcessor(DocumentProcessor):
    """XML文件处理器"""
    
    def check_validity(self, xml_path: Union[str, Path]) -> Optional[str]:
        """检查XML文件有效性"""
        try:
            ET.parse(xml_path)
            with open(xml_path, encoding='utf-8') as f:
                return f.read()
        except ET.ParseError as e:
            logger.error(f"XML格式错误：{e}")
            return None
    
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """加载XML文件"""
        documents = []
        xml_files = glob(os.path.join(file_path, '*.xml')) if Path(file_path).is_dir() else [file_path]
        
        for xml_file in xml_files:
            data = self.check_validity(xml_file)
            if not data:
                continue
                
            sub_name = Path(xml_file).stem
            metadata = {"file_path": str(xml_file)}
            
            if "嵌套" in sub_name:
                parts = sub_name.split("-")
                node1 = parts[0].split("_")[-1]
                node2 = parts[2].split("_")[-1]
                content = f"嵌套: {data}"
                metadata.update({"node1": node1, "node2": node2, "relation": '嵌套'})
            elif "完整" in sub_name:
                node = sub_name.split("_")[1]
                content = f"完整: {data}"
                metadata["node"] = node
            else:
                node = sub_name.split("_")[1] if "_" in sub_name else sub_name
                content = f"单节点: {data}"
                metadata["node"] = node
            
            documents.append(Document(page_content=content, metadata=metadata))
        
        return documents


class ExcelProcessor(DocumentProcessor):
    """Excel文件处理器"""
    
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """加载Excel文件"""
        try:
            df = pd.read_excel(file_path)
            documents = []
            
            # 动态检测列名
            content_cols = [col for col in ['问题', '答案', 'question', 'answer'] if col in df.columns]
            metadata_cols = [col for col in ['序号', '备注', 'id', 'note'] if col in df.columns]
            
            for index, row in df.iterrows():
                # 组合内容列
                content = " ".join(f"{col}: {str(row[col])}" for col in content_cols if pd.notna(row[col]))
                
                # 创建元数据
                metadata = {"file_path": str(file_path), "row_index": index}
                for col in metadata_cols:
                    if col in df.columns and pd.notna(row[col]):
                        metadata[col] = row[col]
                
                if content.strip():  # 只添加非空文档
                    documents.append(Document(page_content=content, metadata=metadata))
            
            return documents
        except Exception as e:
            logger.error(f"解析Excel文件时出错: {str(e)}")
            return []


class TextProcessor(DocumentProcessor):
    """文本文件处理器（Word, Markdown, Text, PDF）"""
    
    def read_word(self, file_path: Union[str, Path]) -> str:
        """读取Word文档"""
        doc = docx.Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return '\n\n'.join(paragraphs)
    
    def read_text(self, file_path: Union[str, Path]) -> str:
        """读取文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def split_text(self, text: str, max_length: int = None, overlap: int = None) -> List[str]:
        """切分文本"""
        max_length = max_length or self.config.chunk_size
        overlap = overlap or self.config.chunk_overlap
        
        if not text or max_length >= len(text):
            return [text] if text else []
        
        # 使用numpy优化切分
        text_length = len(text)
        step = max_length - overlap
        num_segments = max(1, (text_length - overlap + step - 1) // step)
        
        starts = np.arange(num_segments) * step
        ends = np.minimum(starts + max_length, text_length)
        
        return [text[int(start):int(end)] for start, end in zip(starts, ends)]
    
    def load_and_split(self, file_path: Union[str, Path], file_type: str) -> List[Document]:
        """加载并切分文本文件"""
        try:
            # 根据文件类型选择读取方法
            if file_type in ('.docx', '.doc'):
                text = self.read_word(file_path)
            elif file_type in ('.txt', '.md'):
                text = self.read_text(file_path)
            elif file_type == '.pdf':
                loader = PyPDFLoader(str(file_path))
                return loader.load()
            else:
                return []
            
            if not text:
                logger.warning(f"文件 {file_path} 不包含文本内容")
                return []
            
            # 切分文本
            segments = self.split_text(text)
            
            # 创建文档
            documents = [
                Document(
                    page_content=segment,
                    metadata={"segment_index": i, "file_path": str(file_path)}
                )
                for i, segment in enumerate(segments)
            ]
            
            return documents
            
        except Exception as e:
            logger.error(f"处理文件 {file_path} 时出错: {str(e)}")
            return []


class AdvancedSplitter:
    """高级文档切分器"""
    
    def __init__(self, config: SplitConfig = None):
        self.config = config or SplitConfig()
    
    def perception_split(self, documents: List[Document], file_type: str) -> List[Document]:
        """内容感知切分（主要用于Markdown）"""
        if file_type != '.md':
            return self.recursive_split(documents)
        
        # 合并文档内容
        all_text = "\n".join([doc.page_content for doc in documents])
        
        # Markdown标题切分
        headers_to_split_on = [
            ("#", "Header 1"),
            ("##", "Header 2"),
            ("###", "Header 3"),
            ("####", "Header 4"),
            ("#####", "Header 5"),
            ("######", "Header 6"),
        ]
        
        splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
        split_docs = splitter.split_text(all_text)
        
        # 保留原始元数据
        if documents:
            for doc in split_docs:
                doc.metadata.update(documents[0].metadata)
        
        return split_docs
    
    def recursive_split(self, documents: List[Document]) -> List[Document]:
        """递归字符切分"""
        all_text = "\n".join([doc.page_content for doc in documents])
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = splitter.split_text(all_text)
        
        # 保留原始元数据
        metadata = documents[0].metadata if documents else {}
        return [Document(page_content=chunk, metadata=metadata.copy()) for chunk in chunks]
    
    def semantic_split(self, documents: List[Document], embeddings_model: OpenAIEmbeddings) -> List[Document]:
        """基于语义相似度的切分"""
        if not documents:
            return []
        
        logger.info("开始语义切分...")
        
        # 构建向量存储
        vector_store = FAISS.from_documents(documents, embeddings_model)
        
        # 获取所有文档
        doc_ids = list(vector_store.index_to_docstore_id.values())
        docs = [vector_store.docstore.search(doc_id) for doc_id in doc_ids]
        docs = [doc for doc in docs if doc is not None]
        
        if not docs:
            return []
        
        texts = [doc.page_content for doc in docs]
        
        # 批量获取嵌入向量
        embeddings = embeddings_model.embed_documents(texts)
        similarity_matrix = cosine_similarity(embeddings)
        
        # 识别切分点
        split_points = [0]
        current_chunk = 0
        
        for i in range(1, len(similarity_matrix)):
            if similarity_matrix[i, current_chunk] < self.config.semantic_threshold:
                split_points.append(i)
                current_chunk = i
        
        split_points.append(len(similarity_matrix))
        
        # 根据切分点创建新文档
        result_docs = []
        for i in range(len(split_points) - 1):
            start, end = split_points[i], split_points[i + 1]
            content = "\n\n".join([docs[j].page_content for j in range(start, end)])
            
            # 合并元数据
            metadata = {}
            for j in range(start, end):
                metadata.update(docs[j].metadata)
            
            result_docs.append(Document(page_content=content, metadata=metadata))
        
        logger.info(f"语义切分完成，生成 {len(result_docs)} 个语义段落")
        return result_docs
    
    def parent_child_split(
        self, 
        documents: List[Document], 
        embeddings_model: OpenAIEmbeddings
    ) -> Tuple[List[Document], List[Document], FAISS, Dict[str, str]]:
        """父子文档切分"""
        logger.info("开始父子切分...")
        
        # 初始化切分器
        large_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.large_chunk_size,
            chunk_overlap=self.config.large_chunk_overlap
        )
        
        small_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap
        )
        
        larger_documents = []
        smaller_documents = []
        small_to_large_mapping = {}
        
        # 第一步：创建大块文档
        for doc in documents:
            large_chunks = large_splitter.split_text(doc.page_content)
            for chunk in large_chunks:
                larger_documents.append(
                    Document(page_content=chunk, metadata=doc.metadata.copy())
                )
        
        # 创建父向量存储
        parent_vector_store = FAISS.from_documents(larger_documents, embeddings_model)
        parent_ids = list(parent_vector_store.index_to_docstore_id.values())
        
        # 第二步：为每个大块创建小块
        for parent_id in parent_ids:
            parent_doc = parent_vector_store.docstore.search(parent_id)
            if not parent_doc:
                continue
            
            # 切分小块
            small_chunks = small_splitter.split_text(parent_doc.page_content)
            small_docs = [
                Document(page_content=chunk, metadata=parent_doc.metadata.copy())
                for chunk in small_chunks
            ]
            
            # 创建子向量存储并合并
            if small_docs:
                child_store = FAISS.from_documents(small_docs, embeddings_model)
                child_ids = list(child_store.index_to_docstore_id.values())
                
                # 建立映射关系
                for child_id in child_ids:
                    small_to_large_mapping[child_id] = parent_id
                
                # 合并到父存储
                parent_vector_store.merge_from(child_store)
                smaller_documents.extend(small_docs)
        
        logger.info(f"父子切分完成：{len(smaller_documents)} 个小块，{len(larger_documents)} 个大块")
        return smaller_documents, larger_documents, parent_vector_store, small_to_large_mapping


class DocumentManager:
    """文档管理器 - 统一的文档处理入口"""
    
    def __init__(self, embeddings_model: Optional[OpenAIEmbeddings] = None, config: SplitConfig = None):
        self.config = config or SplitConfig()
        self.embeddings_model = embeddings_model or LLM_INIT().create_embeddings_client()
        
        # 初始化处理器
        self.xml_processor = XMLProcessor(self.config)
        self.excel_processor = ExcelProcessor(self.config)
        self.text_processor = TextProcessor(self.config)
        self.splitter = AdvancedSplitter(self.config)
        
        # 文件类型映射
        self.processor_mapping = {
            '.xml': self.xml_processor,
            '.xlsx': self.excel_processor,
            '.xls': self.excel_processor,
            '.csv': self.excel_processor,
            '.md': self.text_processor,
            '.docx': self.text_processor,
            '.doc': self.text_processor,
            '.txt': self.text_processor,
            '.pdf': self.text_processor
        }
    
    def load_document(self, file_path: Union[str, Path]) -> List[Document]:
        """加载文档"""
        path = Path(file_path)
        file_type = path.suffix.lower()
        
        processor = self.processor_mapping.get(file_type)
        if not processor:
            raise ValueError(f"不支持的文件类型: {file_type}")
        
        if isinstance(processor, XMLProcessor):
            return processor.load(file_path)
        elif isinstance(processor, ExcelProcessor):
            return processor.load(file_path)
        elif isinstance(processor, TextProcessor):
            return processor.load_and_split(file_path, file_type)
        
        return []
    
    def process_document(
        self,
        file_path: Union[str, Path],
        split_type: Union[str, SplitType] = SplitType.SIMPLE
    ) -> Union[List[Document], Tuple[List[Document], List[Document], FAISS, Dict[str, str]]]:
        """处理文档的主入口"""
        # 验证文件
        valid, result = self.text_processor.validate_file(file_path)
        if not valid:
            logger.error(f"文件验证失败: {result}")
            return [] if split_type != SplitType.PARENT_CHILD else ([], [], None, {})
        
        file_type = result
        
        # 转换split_type为枚举
        if isinstance(split_type, str):
            try:
                split_type = SplitType(split_type)
            except ValueError:
                logger.error(f"不支持的切分类型: {split_type}")
                return [] if split_type != "parent_child" else ([], [], None, {})
        
        try:
            # 加载文档
            documents = self.load_document(file_path)
            if not documents:
                logger.warning(f"未能从文件加载任何文档: {file_path}")
                return [] if split_type != SplitType.PARENT_CHILD else ([], [], None, {})
            
            logger.info(f"已加载 {len(documents)} 个文档")
            
            # 根据切分类型处理
            if split_type == SplitType.SIMPLE:
                return documents
            
            elif split_type == SplitType.PERCEPTION:
                return self.splitter.perception_split(documents, file_type)
            
            elif split_type == SplitType.RECURSION:
                return self.splitter.recursive_split(documents)
            
            elif split_type == SplitType.SEMANTIC:
                return self.splitter.semantic_split(documents, self.embeddings_model)
            
            elif split_type == SplitType.PARENT_CHILD:
                return self.splitter.parent_child_split(documents, self.embeddings_model)
            
            else:
                return documents
                
        except Exception as e:
            logger.error(f"处理文档时出错: {str(e)}")
            return [] if split_type != SplitType.PARENT_CHILD else ([], [], None, {})


# 保留向后兼容的函数
def check_and_create_dir(dir_path: Union[str, Path]) -> bool:
    """向后兼容函数"""
    processor = DocumentProcessor()
    return processor.check_and_create_dir(dir_path)


def process_document(
    file_path: Union[str, Path],
    embedding_model: OpenAIEmbeddings,
    split_type: str = 'simple',
    **kwargs
) -> Union[List[Document], Tuple[List[Document], List[Document], FAISS, Dict[str, str]]]:
    """向后兼容函数"""
    config = SplitConfig(
        chunk_size=kwargs.get('max_length', 500),
        chunk_overlap=kwargs.get('overlap', 100),
        large_chunk_size=kwargs.get('large_chunk_size', 1000),
        large_chunk_overlap=kwargs.get('large_chunk_overlap', 200),
        semantic_threshold=kwargs.get('chunk_threshold', 0.7)
    )
    
    manager = DocumentManager(embedding_model, config)
    return manager.process_document(file_path, split_type)


if __name__ == '__main__':
    # 测试代码
    test_files = [
        "C:/Users/16000/Desktop/ww.docx",
        "C:/Users/16000/Desktop/ww.md",
        "C:/Users/16000/Desktop/test.xml",
        "C:/Users/16000/Desktop/data.xlsx"
    ]
    
    # 创建文档管理器
    manager = DocumentManager()
    
    for test_file in test_files:
        if Path(test_file).exists():
            print(f"\n测试文件: {test_file}")
            
            # 测试简单加载
            docs = manager.process_document(test_file, SplitType.SIMPLE)
            print(f"简单加载: {len(docs)} 个文档")
            
            # 测试父子切分
            if Path(test_file).suffix.lower() in ['.docx', '.md', '.txt']:
                result = manager.process_document(test_file, SplitType.PARENT_CHILD)
                if isinstance(result, tuple):
                    small, large, store, mapping = result
                    print(f"父子切分: {len(small)} 个小块, {len(large)} 个大块")